#!/usr/bin/env python3
"""
Smart Ingestion System for Lumen Navigator RAG

A comprehensive document processing and ingestion system designed for children's home management
with focus on regulatory compliance, safeguarding, and policy management.

Features:
- Intelligent document classification and taxonomy
- Selective index population based on document type
- Performance-optimized chunking strategies
- Regulatory compliance tagging
- Automated metadata extraction
- Document freshness management
- Integration with existing smart query router
"""

import os
import json
import time
import hashlib
import logging
from typing import Dict, List, Any, Optional, Tuple, Union
from dataclasses import dataclass, asdict
from datetime import datetime, timedelta
from pathlib import Path
import shutil
import re

# Core libraries
import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter, TokenTextSplitter
from langchain.docstore.document import Document
from langchain_openai import OpenAIEmbeddings
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import FAISS
import tiktoken

# Document processing
import PyPDF2
import docx
from bs4 import BeautifulSoup
import pandas as pd
import fitz  # PyMuPDF for better PDF processing

# Advanced processing
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer
from collections import Counter

# Integration with existing system
from smart_query_router import PerformanceTracker

# =============================================================================
# CONFIGURATION
# =============================================================================

# Directory structure (matching existing smart router)
DB_BASE_DIR = "indexes"
OPENAI_DB_DIR = os.path.join(DB_BASE_DIR, "openai_index")
GOOGLE_DB_DIR = os.path.join(DB_BASE_DIR, "google_index")
METADATA_DIR = "metadata"
PERFORMANCE_DIR = "performance_metrics"
DOCUMENTS_DIR = "documents"
PROCESSING_DIR = "document_processing"

# Document taxonomy for children's homes
DOCUMENT_TAXONOMY = {
    "regulatory_compliance": {
        "ofsted_requirements": ["ofsted", "inspection", "standards", "regulation", "requirement"],
        "government_regulations": ["legislation", "act", "statutory", "legal", "government", "department"],
        "local_authority": ["local authority", "council", "commissioning", "placement", "contract"]
    },
    "safeguarding": {
        "safeguarding_policies": ["safeguarding", "child protection", "abuse", "neglect", "welfare"],
        "risk_assessments": ["risk assessment", "hazard", "safety", "evaluation", "mitigation"],
        "incident_procedures": ["incident", "emergency", "crisis", "accident", "reporting"]
    },
    "operational": {
        "daily_operations": ["routine", "schedule", "daily", "operations", "activities"],
        "health_safety": ["health", "safety", "medical", "hygiene", "infection", "medication"],
        "staff_procedures": ["staff", "personnel", "training", "supervision", "performance"]
    },
    "care_support": {
        "care_planning": ["care plan", "assessment", "goals", "outcomes", "review"],
        "educational_support": ["education", "school", "learning", "academic", "tuition"],
        "therapeutic": ["therapy", "counseling", "mental health", "psychological", "wellbeing"]
    },
    "administration": {
        "hr_policies": ["human resources", "employment", "recruitment", "disciplinary"],
        "financial": ["finance", "budget", "expenditure", "procurement", "accounts"],
        "records": ["records", "documentation", "filing", "data protection", "confidentiality"]
    }
}

# Index routing strategy - which provider handles which document types
INDEX_ROUTING = {
    "openai_primary": [
        "operational", "care_support", "administration"
    ],
    "google_primary": [
        "regulatory_compliance", "safeguarding"
    ],
    "both_indexes": [
        "safeguarding.safeguarding_policies",
        "safeguarding.incident_procedures",
        "regulatory_compliance.ofsted_requirements"
    ]
}

# Chunking strategies per provider
CHUNKING_STRATEGIES = {
    "openai": {
        "chunk_size": 1000,
        "chunk_overlap": 200,
        "separators": ["\n\n", "\n", ". ", " ", ""],
        "length_function": "tiktoken"
    },
    "google": {
        "chunk_size": 1200,
        "chunk_overlap": 150,
        "separators": ["\n\n", "\n", ". ", " ", ""],
        "length_function": "len"
    }
}

# Regulatory frameworks and their importance
REGULATORY_FRAMEWORKS = {
    "ofsted": {
        "priority": "critical",
        "keywords": ["ofsted", "inspection", "standards", "quality", "safeguarding children"],
        "compliance_weight": 1.0
    },
    "government": {
        "priority": "critical", 
        "keywords": ["children act", "care standards", "statutory", "regulation", "legislation"],
        "compliance_weight": 1.0
    },
    "local_authority": {
        "priority": "high",
        "keywords": ["local authority", "commissioning", "placement", "social services"],
        "compliance_weight": 0.8
    },
    "internal": {
        "priority": "medium",
        "keywords": ["policy", "procedure", "guideline", "internal"],
        "compliance_weight": 0.6
    }
}

# Quality thresholds
QUALITY_THRESHOLDS = {
    "min_chunk_length": 100,
    "max_chunk_length": 2000,
    "min_document_length": 50,
    "duplicate_threshold": 0.85
}

# Logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# =============================================================================
# DATA STRUCTURES
# =============================================================================

@dataclass
class DocumentMetadata:
    """Enhanced metadata structure for children's home documents"""
    document_id: str
    file_path: str
    file_name: str
    file_size_bytes: int
    file_type: str
    creation_date: str
    modification_date: str
    processing_date: str
    
    # Classification
    document_category: str
    document_subcategory: str
    regulatory_framework: str
    compliance_criticality: float
    authority_level: str
    
    # Content analysis
    content_hash: str
    total_chunks: int
    average_chunk_length: int
    language: str
    readability_score: float
    
    # Performance tracking
    index_assignments: List[str]
    processing_time_seconds: float
    embedding_model_used: str
    
    # Version management
    version: str
    supersedes: Optional[str]
    is_current: bool

@dataclass
class ChunkMetadata:
    """Metadata for individual chunks"""
    chunk_id: str
    document_id: str
    chunk_index: int
    chunk_text: str
    chunk_length: int
    
    # Context
    section_title: str
    page_number: Optional[int]
    paragraph_index: int
    
    # Classification
    content_type: str
    importance_score: float
    regulatory_relevance: float
    
    # Processing
    embedding_provider: str
    chunk_creation_date: str

# =============================================================================
# DOCUMENT CLASSIFIER
# =============================================================================

class DocumentClassifier:
    """Intelligent document classification for children's home content"""
    
    def __init__(self):
        self.taxonomy = DOCUMENT_TAXONOMY
        self.regulatory_frameworks = REGULATORY_FRAMEWORKS
        self.vectorizer = TfidfVectorizer(max_features=1000, stop_words='english')
        self._load_nlp_model()
    
    def _load_nlp_model(self):
        """Load spaCy model for advanced text processing"""
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            logger.warning("spaCy model not found. Install with: python -m spacy download en_core_web_sm")
            self.nlp = None
    
    def classify_document(self, content: str, file_name: str) -> Dict[str, Any]:
        """
        Classify document into taxonomy categories and extract metadata
        
        Args:
            content (str): Document text content
            file_name (str): Original file name
            
        Returns:
            Dict[str, Any]: Classification results and metadata
        """
        content_lower = content.lower()
        file_name_lower = file_name.lower()
        
        # Calculate category scores
        category_scores = {}
        for category, subcategories in self.taxonomy.items():
            category_scores[category] = {}
            category_total = 0
            
            for subcategory, keywords in subcategories.items():
                score = self._calculate_keyword_score(content_lower, keywords)
                # Boost score if keywords appear in filename
                if any(keyword in file_name_lower for keyword in keywords):
                    score *= 1.5
                
                category_scores[category][subcategory] = score
                category_total += score
            
            category_scores[category]["total"] = category_total
        
        # Determine best category and subcategory
        best_category = max(category_scores.items(), key=lambda x: x[1]["total"])
        category_name = best_category[0]
        subcategories = best_category[1]
        
        # Find best subcategory (excluding 'total')
        subcategory_scores = {k: v for k, v in subcategories.items() if k != "total"}
        best_subcategory = max(subcategory_scores.items(), key=lambda x: x[1]) if subcategory_scores else ("unknown", 0)
        
        # Determine regulatory framework
        regulatory_info = self._identify_regulatory_framework(content_lower, file_name_lower)
        
        # Calculate compliance criticality
        compliance_criticality = self._calculate_compliance_criticality(
            category_name, best_subcategory[0], regulatory_info
        )
        
        # Extract additional metadata
        additional_metadata = self._extract_additional_metadata(content)
        
        return {
            "category": category_name,
            "subcategory": best_subcategory[0],
            "confidence": min(best_subcategory[1], 1.0),
            "regulatory_framework": regulatory_info["framework"],
            "authority_level": regulatory_info["authority_level"],
            "compliance_criticality": compliance_criticality,
            "all_scores": category_scores,
            **additional_metadata
        }
    
    def _calculate_keyword_score(self, content: str, keywords: List[str]) -> float:
        """Calculate relevance score based on keyword presence"""
        score = 0.0
        content_words = content.split()
        total_words = len(content_words)
        
        if total_words == 0:
            return 0.0
        
        for keyword in keywords:
            # Exact phrase matching
            if keyword in content:
                score += content.count(keyword) * 2.0
            
            # Individual word matching
            keyword_words = keyword.split()
            for word in keyword_words:
                if word in content_words:
                    score += content_words.count(word) * 0.5
        
        # Normalize by content length
        return min(score / total_words * 100, 1.0)
    
    def _identify_regulatory_framework(self, content: str, file_name: str) -> Dict[str, str]:
        """Identify the regulatory framework the document belongs to"""
        framework_scores = {}
        
        for framework, info in self.regulatory_frameworks.items():
            score = 0.0
            for keyword in info["keywords"]:
                if keyword in content:
                    score += content.count(keyword) * info["compliance_weight"]
                if keyword in file_name:
                    score += 2.0 * info["compliance_weight"]
            
            framework_scores[framework] = score
        
        if not framework_scores or max(framework_scores.values()) == 0:
            return {"framework": "internal", "authority_level": "internal"}
        
        best_framework = max(framework_scores.items(), key=lambda x: x[1])
        framework_name = best_framework[0]
        
        # Determine authority level
        authority_mapping = {
            "ofsted": "national_regulator",
            "government": "statutory", 
            "local_authority": "local_regulator",
            "internal": "internal"
        }
        
        return {
            "framework": framework_name,
            "authority_level": authority_mapping.get(framework_name, "internal")
        }
    
    def _calculate_compliance_criticality(self, category: str, subcategory: str, 
                                        regulatory_info: Dict[str, str]) -> float:
        """Calculate compliance criticality score (0.0 to 1.0)"""
        base_scores = {
            "regulatory_compliance": 1.0,
            "safeguarding": 0.95,
            "operational": 0.7,
            "care_support": 0.8,
            "administration": 0.5
        }
        
        category_score = base_scores.get(category, 0.5)
        
        # Boost for critical subcategories
        critical_subcategories = {
            "safeguarding_policies": 0.1,
            "incident_procedures": 0.1,
            "ofsted_requirements": 0.15,
            "risk_assessments": 0.05
        }
        
        subcategory_boost = critical_subcategories.get(subcategory, 0.0)
        
        # Framework weight
        framework_weights = {
            "ofsted": 1.0,
            "government": 1.0,
            "local_authority": 0.8,
            "internal": 0.6
        }
        
        framework_weight = framework_weights.get(regulatory_info["framework"], 0.6)
        
        return min((category_score + subcategory_boost) * framework_weight, 1.0)
    
    def _extract_additional_metadata(self, content: str) -> Dict[str, Any]:
        """Extract additional metadata from content"""
        metadata = {
            "language": "en",  # Default to English
            "readability_score": 0.5,  # Placeholder
            "content_length": len(content),
            "word_count": len(content.split()),
            "sentence_count": len(re.split(r'[.!?]+', content))
        }
        
        if self.nlp:
            try:
                doc = self.nlp(content[:10000])  # Limit for performance
                
                # Extract entities
                entities = [(ent.text, ent.label_) for ent in doc.ents]
                metadata["entities"] = entities[:20]  # Limit entities
                
                # Extract key phrases (noun phrases)
                key_phrases = [chunk.text for chunk in doc.noun_chunks]
                metadata["key_phrases"] = key_phrases[:15]
                
            except Exception as e:
                logger.warning(f"NLP processing failed: {e}")
        
        return metadata

# =============================================================================
# DOCUMENT PROCESSOR
# =============================================================================

class DocumentProcessor:
    """Handles document loading and text extraction"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_doc,
            '.txt': self._process_txt,
            '.html': self._process_html,
            '.md': self._process_markdown
        }
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process document and extract text with metadata
        
        Args:
            file_path (str): Path to document file
            
        Returns:
            Dict[str, Any]: Processed content and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Get file metadata
        file_stats = file_path.stat()
        file_metadata = {
            "file_path": str(file_path),
            "file_name": file_path.name,
            "file_size_bytes": file_stats.st_size,
            "file_type": file_extension,
            "creation_date": datetime.fromtimestamp(file_stats.st_ctime).isoformat(),
            "modification_date": datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
            "processing_date": datetime.now().isoformat()
        }
        
        # Process content
        try:
            content_result = self.supported_formats[file_extension](file_path)
            
            # Validate content
            if not content_result["content"] or len(content_result["content"].strip()) < QUALITY_THRESHOLDS["min_document_length"]:
                raise ValueError("Document content is too short or empty")
            
            return {
                **file_metadata,
                **content_result,
                "processing_success": True,
                "error": None
            }
            
        except Exception as e:
            logger.error(f"Failed to process {file_path}: {e}")
            return {
                **file_metadata,
                "content": "",
                "page_content": [],
                "processing_success": False,
                "error": str(e)
            }
    
    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Process PDF using PyMuPDF for better extraction"""
        try:
            doc = fitz.open(str(file_path))
            content = ""
            page_content = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                
                if page_text.strip():
                    content += page_text + "\n\n"
                    page_content.append({
                        "page_number": page_num + 1,
                        "content": page_text,
                        "word_count": len(page_text.split())
                    })
            
            doc.close()
            
            return {
                "content": content.strip(),
                "page_content": page_content,
                "total_pages": len(doc),
                "extraction_method": "pymupdf"
            }
            
        except Exception as e:
            # Fallback to PyPDF2
            logger.warning(f"PyMuPDF failed, trying PyPDF2: {e}")
            return self._process_pdf_fallback(file_path)
    
    def _process_pdf_fallback(self, file_path: Path) -> Dict[str, Any]:
        """Fallback PDF processing with PyPDF2"""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            content = ""
            page_content = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                
                if page_text.strip():
                    content += page_text + "\n\n"
                    page_content.append({
                        "page_number": page_num + 1,
                        "content": page_text,
                        "word_count": len(page_text.split())
                    })
            
            return {
                "content": content.strip(),
                "page_content": page_content,
                "total_pages": len(pdf_reader.pages),
                "extraction_method": "pypdf2"
            }
    
    def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Process Word document"""
        doc = docx.Document(str(file_path))
        content = ""
        paragraphs = []
        
        for para_num, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                content += paragraph.text + "\n\n"
                paragraphs.append({
                    "paragraph_number": para_num + 1,
                    "content": paragraph.text,
                    "style": paragraph.style.name if paragraph.style else "Normal"
                })
        
        return {
            "content": content.strip(),
            "page_content": paragraphs,
            "total_paragraphs": len(paragraphs),
            "extraction_method": "python-docx"
        }
    
    def _process_doc(self, file_path: Path) -> Dict[str, Any]:
        """Process legacy Word document"""
        # This would require additional libraries like python-docx2txt
        # For now, return error asking for conversion
        raise ValueError("Legacy .doc files not supported. Please convert to .docx format.")
    
    def _process_txt(self, file_path: Path) -> Dict[str, Any]:
        """Process plain text file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            content = file.read()
        
        lines = content.split('\n')
        
        return {
            "content": content,
            "page_content": [{"line_number": i+1, "content": line} for i, line in enumerate(lines) if line.strip()],
            "total_lines": len(lines),
            "extraction_method": "text"
        }
    
    def _process_html(self, file_path: Path) -> Dict[str, Any]:
        """Process HTML file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            html_content = file.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        content = soup.get_text()
        
        # Clean up whitespace
        lines = (line.strip() for line in content.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        content = ' '.join(chunk for chunk in chunks if chunk)
        
        return {
            "content": content,
            "page_content": [{"section": "html_body", "content": content}],
            "extraction_method": "beautifulsoup"
        }
    
    def _process_markdown(self, file_path: Path) -> Dict[str, Any]:
        """Process Markdown file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            content = file.read()
        
        # Simple markdown processing - could be enhanced with markdown library
        return {
            "content": content,
            "page_content": [{"section": "markdown", "content": content}],
            "extraction_method": "text"
        }

# =============================================================================
# SMART CHUNKING ENGINE
# =============================================================================

class SmartChunkingEngine:
    """Advanced chunking with provider-specific strategies"""
    
    def __init__(self):
        self.strategies = CHUNKING_STRATEGIES
        self.tiktoken_encoding = tiktoken.get_encoding("cl100k_base")
    
    def create_chunks(self, content: str, provider: str, document_metadata: Dict[str, Any]) -> List[Document]:
        """
        Create optimized chunks for specific embedding provider
        
        Args:
            content (str): Document content
            provider (str): Target embedding provider (openai/google)
            document_metadata (Dict[str, Any]): Document metadata
            
        Returns:
            List[Document]: Chunked documents with metadata
        """
        if provider not in self.strategies:
            raise ValueError(f"Unknown provider: {provider}")
        
        strategy = self.strategies[provider]
        
        # Select appropriate text splitter
        if strategy["length_function"] == "tiktoken":
            text_splitter = TokenTextSplitter(
                chunk_size=strategy["chunk_size"],
                chunk_overlap=strategy["chunk_overlap"],
                encoding_name="cl100k_base"
            )
        else:
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=strategy["chunk_size"],
                chunk_overlap=strategy["chunk_overlap"],
                separators=strategy["separators"],
                length_function=len
            )
        
        # Create base chunks
        base_chunks = text_splitter.split_text(content)
        
        # Apply intelligent chunking enhancements
        enhanced_chunks = self._enhance_chunks(base_chunks, document_metadata, provider)
        
        # Convert to Document objects with metadata
        documents = []
        for i, chunk in enumerate(enhanced_chunks):
            if len(chunk.strip()) >= QUALITY_THRESHOLDS["min_chunk_length"]:
                
                chunk_metadata = {
                    "chunk_id": f"{document_metadata.get('document_id', 'unknown')}_{provider}_{i}",
                    "document_id": document_metadata.get('document_id'),
                    "chunk_index": i,
                    "total_chunks": len(enhanced_chunks),
                    "chunk_length": len(chunk),
                    "provider": provider,
                    "document_category": document_metadata.get('category'),
                    "document_subcategory": document_metadata.get('subcategory'),
                    "regulatory_framework": document_metadata.get('regulatory_framework'),
                    "compliance_criticality": document_metadata.get('compliance_criticality'),
                    "file_name": document_metadata.get('file_name'),
                    "creation_date": datetime.now().isoformat()
                }
                
                # Add provider-specific metadata
                if provider == "openai":
                    chunk_metadata["token_count"] = len(self.tiktoken_encoding.encode(chunk))
                
                doc = Document(
                    page_content=chunk,
                    metadata=chunk_metadata
                )
                documents.append(doc)
        
        return documents
    
    def _enhance_chunks(self, chunks: List[str], document_metadata: Dict[str, Any], provider: str) -> List[str]:
        """Apply intelligent enhancements to chunks"""
        enhanced_chunks = []
        
        for chunk in chunks:
            # Skip very short chunks
            if len(chunk.strip()) < QUALITY_THRESHOLDS["min_chunk_length"]:
                continue
            
            # Apply content-aware processing
            processed_chunk = self._process_chunk_content(chunk, document_metadata)
            
            # Apply provider-specific optimizations
            if provider == "openai":
                processed_chunk = self._optimize_for_openai(processed_chunk, document_metadata)
            elif provider == "google":
                processed_chunk = self._optimize_for_google(processed_chunk, document_metadata)
            
            enhanced_chunks.append(processed_chunk)
        
        return enhanced_chunks
    
    def _process_chunk_content(self, chunk: str, metadata: Dict[str, Any]) -> str:
        """Apply content-aware processing to chunks"""
        # Clean up chunk
        chunk = re.sub(r'\s+', ' ', chunk)  # Normalize whitespace
        chunk = chunk.strip()
        
        # Add context for regulatory content
        if metadata.get('category') == 'regulatory_compliance':
            # Ensure regulatory context is preserved
            if not any(word in chunk.lower() for word in ['regulation', 'requirement', 'standard', 'compliance']):
                context = f"[Regulatory Document: {metadata.get('file_name', 'Unknown')}] "
                chunk = context + chunk
        
        # Add context for safeguarding content
        elif metadata.get('category') == 'safeguarding':
            if not any(word in chunk.lower() for word in ['safeguarding', 'child protection', 'safety']):
                context = f"[Safeguarding Policy: {metadata.get('file_name', 'Unknown')}] "
                chunk = context + chunk
        
        return chunk
    
    def _optimize_for_openai(self, chunk: str, metadata: Dict[str, Any]) -> str:
        """Apply OpenAI-specific optimizations"""
        # OpenAI models prefer structured, clear content
        # Add clear section markers for better understanding
        if metadata.get('category') in ['operational', 'care_support']:
            return chunk
        
        return chunk
    
    def _optimize_for_google(self, chunk: str, metadata: Dict[str, Any]) -> str:
        """Apply Google-specific optimizations"""
        # Google models handle regulatory/legal content well
        # Preserve formal language and structure
        if metadata.get('category') in ['regulatory_compliance', 'safeguarding']:
            return chunk
        
        return chunk

# =============================================================================
# INDEX MANAGER
# =============================================================================

class IndexManager:
    """Manages FAISS indexes and embedding operations"""
    
    def __init__(self):
        self.embedding_models = {}
        self.vector_stores = {}
        self.performance_tracker = PerformanceTracker()
        self._ensure_directories()
        self._load_embedding_models()
        self._load_existing_indexes()
    
    def _ensure_directories(self):
        """Create necessary directories"""
        for directory in [DB_BASE_DIR, METADATA_DIR, PERFORMANCE_DIR, PROCESSING_DIR]:
            Path(directory).mkdir(exist_ok=True)
    
    def _load_embedding_models(self):
        """Load embedding models"""
        try:
            # Use the same model as your query router (text-embedding-3-large)
            self.embedding_models["openai"] = OpenAIEmbeddings(
                model="text-embedding-3-large",
                show_progress_bar=False
            )
            logger.info("Loaded OpenAI embedding model")
        except Exception as e:
            logger.error(f"Failed to load OpenAI embedding model: {e}")
        
        try:
            self.embedding_models["google"] = GoogleGenerativeAIEmbeddings(
                model="models/embedding-001"
            )
            logger.info("Loaded Google embedding model")
        except Exception as e:
            logger.error(f"Failed to load Google embedding model: {e}")
    
    def _load_existing_indexes(self):
        """Load existing FAISS indexes"""
        index_paths = {
            "openai": OPENAI_DB_DIR,
            "google": GOOGLE_DB_DIR
        }
        
        for provider, path in index_paths.items():
            if os.path.exists(path) and provider in self.embedding_models:
                try:
                    self.vector_stores[provider] = FAISS.load_local(
                        path,
                        self.embedding_models[provider],
                        allow_dangerous_deserialization=True
                    )
                    logger.info(f"Loaded existing {provider} index")
                except Exception as e:
                    logger.warning(f"Could not load existing {provider} index: {e}")
    
    def add_documents(self, documents: List[Document], provider: str) -> bool:
        """
        Add documents to specific provider index
        
        Args:
            documents (List[Document]): Documents to add
            provider (str): Target provider (openai/google)
            
        Returns:
            bool: Success status
        """
        if provider not in self.embedding_models:
            logger.error(f"Provider {provider} not available")
            return False
        
        try:
            start_time = time.time()
            
            if provider in self.vector_stores:
                # Add to existing index
                self.vector_stores[provider].add_documents(documents)
            else:
                # Create new index
                self.vector_stores[provider] = FAISS.from_documents(
                    documents,
                    self.embedding_models[provider]
                )
            
            # Save index
            index_path = OPENAI_DB_DIR if provider == "openai" else GOOGLE_DB_DIR
            self.vector_stores[provider].save_local(index_path)
            
            processing_time = time.time() - start_time
            
            # Record performance metrics
            self.performance_tracker.record_query(
                provider=provider,
                query=f"ingestion_{len(documents)}_docs",
                response_time=processing_time,
                result_count=len(documents),
                pattern="ingestion",
                success=True
            )
            
            logger.info(f"Successfully added {len(documents)} documents to {provider} index in {processing_time:.2f}s")
            return True
            
        except Exception as e:
            logger.error(f"Failed to add documents to {provider} index: {e}")
            return False
    
    def get_index_stats(self) -> Dict[str, Any]:
        """Get statistics for all indexes"""
        stats = {}
        
        for provider, vector_store in self.vector_stores.items():
            try:
                # Get document count
                doc_count = vector_store.index.ntotal if hasattr(vector_store.index, 'ntotal') else 0
                
                stats[provider] = {
                    "document_count": doc_count,
                    "index_path": OPENAI_DB_DIR if provider == "openai" else GOOGLE_DB_DIR,
                    "status": "active"
                }
            except Exception as e:
                stats[provider] = {
                    "document_count": 0,
                    "error": str(e),
                    "status": "error"
                }
        
        return stats

# =============================================================================
# DOCUMENT DEDUPLICATOR
# =============================================================================

class DocumentDeduplicator:
    """Handles document deduplication and version management"""
    
    def __init__(self):
        self.processed_hashes = self._load_processed_hashes()
    
    def _load_processed_hashes(self) -> Dict[str, Dict[str, Any]]:
        """Load previously processed document hashes"""
        hash_file = Path(METADATA_DIR) / "document_hashes.json"
        if hash_file.exists():
            try:
                with open(hash_file, 'r') as f:
                    return json.load(f)
            except Exception as e:
                logger.warning(f"Could not load document hashes: {e}")
        return {}
    
    def _save_processed_hashes(self):
        """Save document hashes to file"""
        hash_file = Path(METADATA_DIR) / "document_hashes.json"
        try:
            with open(hash_file, 'w') as f:
                json.dump(self.processed_hashes, f, indent=2)
        except Exception as e:
            logger.error(f"Could not save document hashes: {e}")
    
    def calculate_content_hash(self, content: str) -> str:
        """Calculate hash for document content"""
        # Normalize content for hashing
        normalized_content = re.sub(r'\s+', ' ', content.lower().strip())
        return hashlib.sha256(normalized_content.encode()).hexdigest()
    
    def is_duplicate(self, content: str, file_path: str) -> Tuple[bool, Optional[str]]:
        """
        Check if document is a duplicate
        
        Args:
            content (str): Document content
            file_path (str): File path
            
        Returns:
            Tuple[bool, Optional[str]]: (is_duplicate, existing_document_id)
        """
        content_hash = self.calculate_content_hash(content)
        
        # Check for exact content match
        for doc_id, info in self.processed_hashes.items():
            if info["content_hash"] == content_hash:
                return True, doc_id
        
        # Check for near-duplicate content using similarity
        if len(self.processed_hashes) > 0:
            similarity_threshold = QUALITY_THRESHOLDS["duplicate_threshold"]
            for doc_id, info in self.processed_hashes.items():
                if self._calculate_similarity(content, info.get("sample_content", "")) > similarity_threshold:
                    return True, doc_id
        
        return False, None
    
    def _calculate_similarity(self, content1: str, content2: str) -> float:
        """Calculate similarity between two documents"""
        if not content1 or not content2:
            return 0.0
        
        # Simple similarity based on common words
        words1 = set(content1.lower().split())
        words2 = set(content2.lower().split())
        
        if not words1 or not words2:
            return 0.0
        
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union)
    
    def register_document(self, document_id: str, content: str, file_path: str, metadata: Dict[str, Any]):
        """Register processed document to prevent future duplicates"""
        content_hash = self.calculate_content_hash(content)
        
        self.processed_hashes[document_id] = {
            "content_hash": content_hash,
            "file_path": file_path,
            "processing_date": datetime.now().isoformat(),
            "sample_content": content[:500],  # Store sample for similarity checking
            "metadata": {
                "file_name": metadata.get("file_name"),
                "file_size": metadata.get("file_size_bytes"),
                "category": metadata.get("category")
            }
        }
        
        self._save_processed_hashes()

# =============================================================================
# MAIN SMART INGESTION ENGINE
# =============================================================================

class SmartIngestionEngine:
    """Main ingestion engine that coordinates all components"""
    
    def __init__(self):
        self.classifier = DocumentClassifier()
        self.processor = DocumentProcessor()
        self.chunking_engine = SmartChunkingEngine()
        self.index_manager = IndexManager()
        self.deduplicator = DocumentDeduplicator()
        
        # Statistics tracking
        self.ingestion_stats = {
            "total_processed": 0,
            "successful_ingestions": 0,
            "failed_ingestions": 0,
            "duplicates_found": 0,
            "start_time": datetime.now().isoformat()
        }
    
    def ingest_document(self, file_path: str, force_reprocess: bool = False) -> Dict[str, Any]:
        """
        Process and ingest a single document
        
        Args:
            file_path (str): Path to document file
            force_reprocess (bool): Force reprocessing even if duplicate
            
        Returns:
            Dict[str, Any]: Ingestion results
        """
        start_time = time.time()
        
        try:
            # Step 1: Process document
            logger.info(f"Processing document: {file_path}")
            doc_result = self.processor.process_document(file_path)
            
            if not doc_result["processing_success"]:
                return self._create_failure_result(file_path, doc_result["error"], start_time)
            
            content = doc_result["content"]
            
            # Step 2: Check for duplicates
            if not force_reprocess:
                is_duplicate, existing_id = self.deduplicator.is_duplicate(content, file_path)
                if is_duplicate:
                    self.ingestion_stats["duplicates_found"] += 1
                    return self._create_duplicate_result(file_path, existing_id, start_time)
            
            # Step 3: Classify document
            logger.info(f"Classifying document: {Path(file_path).name}")
            classification = self.classifier.classify_document(content, doc_result["file_name"])
            
            # Step 4: Create document metadata
            document_id = self._generate_document_id(doc_result["file_name"], classification)
            
            document_metadata = DocumentMetadata(
                document_id=document_id,
                file_path=doc_result["file_path"],
                file_name=doc_result["file_name"],
                file_size_bytes=doc_result["file_size_bytes"],
                file_type=doc_result["file_type"],
                creation_date=doc_result["creation_date"],
                modification_date=doc_result["modification_date"],
                processing_date=doc_result["processing_date"],
                
                # Classification
                document_category=classification["category"],
                document_subcategory=classification["subcategory"],
                regulatory_framework=classification["regulatory_framework"],
                compliance_criticality=classification["compliance_criticality"],
                authority_level=classification["authority_level"],
                
                # Content analysis
                content_hash=self.deduplicator.calculate_content_hash(content),
                total_chunks=0,  # Will be updated after chunking
                average_chunk_length=0,  # Will be updated after chunking
                language=classification.get("language", "en"),
                readability_score=classification.get("readability_score", 0.5),
                
                # Performance tracking
                index_assignments=[],  # Will be updated based on routing
                processing_time_seconds=0,  # Will be updated at end
                embedding_model_used="",  # Will be updated based on routing
                
                # Version management
                version="1.0",
                supersedes=None,
                is_current=True
            )
            
            # Step 5: Determine index routing
            routing_decision = self._determine_index_routing(classification)
            document_metadata.index_assignments = routing_decision["indexes"]
            
            # Step 6: Process chunks for each assigned index
            all_chunks_created = 0
            total_chunk_length = 0
            
            ingestion_results = {}
            
            for provider in routing_decision["indexes"]:
                logger.info(f"Creating chunks for {provider} index")
                
                # Create provider-specific chunks
                chunks = self.chunking_engine.create_chunks(
                    content, provider, asdict(document_metadata)
                )
                
                if chunks:
                    # Add to index
                    success = self.index_manager.add_documents(chunks, provider)
                    
                    ingestion_results[provider] = {
                        "success": success,
                        "chunk_count": len(chunks),
                        "avg_chunk_length": sum(len(c.page_content) for c in chunks) / len(chunks)
                    }
                    
                    if success:
                        all_chunks_created += len(chunks)
                        total_chunk_length += sum(len(c.page_content) for c in chunks)
                else:
                    ingestion_results[provider] = {
                        "success": False,
                        "error": "No valid chunks created",
                        "chunk_count": 0
                    }
            
            # Update metadata with final statistics
            document_metadata.total_chunks = all_chunks_created
            document_metadata.average_chunk_length = int(total_chunk_length / all_chunks_created) if all_chunks_created > 0 else 0
            document_metadata.processing_time_seconds = time.time() - start_time
            document_metadata.embedding_model_used = ",".join(routing_decision["indexes"])
            
            # Step 7: Save metadata and register document
            self._save_document_metadata(document_metadata)
            self.deduplicator.register_document(
                document_id, content, file_path, asdict(document_metadata)
            )
            
            # Update statistics
            self.ingestion_stats["total_processed"] += 1
            if any(result["success"] for result in ingestion_results.values()):
                self.ingestion_stats["successful_ingestions"] += 1
            else:
                self.ingestion_stats["failed_ingestions"] += 1
            
            return {
                "success": True,
                "document_id": document_id,
                "file_path": file_path,
                "classification": classification,
                "routing": routing_decision,
                "ingestion_results": ingestion_results,
                "metadata": asdict(document_metadata),
                "processing_time": time.time() - start_time,
                "timestamp": datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Ingestion failed for {file_path}: {e}")
            self.ingestion_stats["failed_ingestions"] += 1
            return self._create_failure_result(file_path, str(e), start_time)
    
    def ingest_directory(self, directory_path: str, recursive: bool = True, 
                        file_patterns: List[str] = None) -> Dict[str, Any]:
        """
        Ingest all documents in a directory
        
        Args:
            directory_path (str): Path to directory
            recursive (bool): Process subdirectories
            file_patterns (List[str]): File patterns to match (e.g., ['*.pdf', '*.docx'])
            
        Returns:
            Dict[str, Any]: Batch ingestion results
        """
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        # Find files to process
        files_to_process = []
        
        if file_patterns is None:
            file_patterns = ['*.pdf', '*.docx', '*.txt', '*.html', '*.md']
        
        for pattern in file_patterns:
            if recursive:
                files_to_process.extend(directory_path.rglob(pattern))
            else:
                files_to_process.extend(directory_path.glob(pattern))
        
        logger.info(f"Found {len(files_to_process)} files to process in {directory_path}")
        
        # Process files
        batch_results = {
            "directory": str(directory_path),
            "total_files": len(files_to_process),
            "results": [],
            "summary": {
                "successful": 0,
                "failed": 0,
                "duplicates": 0,
                "total_processing_time": 0
            },
            "start_time": datetime.now().isoformat()
        }
        
        for file_path in files_to_process:
            try:
                result = self.ingest_document(str(file_path))
                batch_results["results"].append(result)
                
                # Update summary
                if result["success"]:
                    if "duplicate" in result:
                        batch_results["summary"]["duplicates"] += 1
                    else:
                        batch_results["summary"]["successful"] += 1
                else:
                    batch_results["summary"]["failed"] += 1
                
                batch_results["summary"]["total_processing_time"] += result.get("processing_time", 0)
                
            except Exception as e:
                logger.error(f"Failed to process {file_path}: {e}")
                batch_results["summary"]["failed"] += 1
                batch_results["results"].append({
                    "success": False,
                    "file_path": str(file_path),
                    "error": str(e),
                    "timestamp": datetime.now().isoformat()
                })
        
        batch_results["end_time"] = datetime.now().isoformat()
        
        # Save batch results
        self._save_batch_results(batch_results)
        
        return batch_results
    
    def _determine_index_routing(self, classification: Dict[str, Any]) -> Dict[str, Any]:
        """Determine which indexes should receive the document"""
        category = classification["category"]
        subcategory = classification["subcategory"]
        full_subcategory = f"{category}.{subcategory}"
        
        # Check for documents that should go to both indexes
        if full_subcategory in INDEX_ROUTING["both_indexes"]:
            return {
                "indexes": ["openai", "google"],
                "reason": f"Critical document type: {full_subcategory}",
                "strategy": "redundancy"
            }
        
        # Check provider-specific routing
        if category in INDEX_ROUTING["openai_primary"]:
            return {
                "indexes": ["openai"],
                "reason": f"OpenAI optimal for category: {category}",
                "strategy": "provider_specialization"
            }
        elif category in INDEX_ROUTING["google_primary"]:
            return {
                "indexes": ["google"],
                "reason": f"Google optimal for category: {category}",
                "strategy": "provider_specialization"
            }
        
        # Fallback: use performance-based routing
        best_provider = self.index_manager.performance_tracker.get_best_provider(category)
        
        return {
            "indexes": [best_provider],
            "reason": f"Performance-based routing for category: {category}",
            "strategy": "performance_optimization"
        }
    
    def _generate_document_id(self, file_name: str, classification: Dict[str, Any]) -> str:
        """Generate unique document ID"""
        # Create ID based on filename, category, and timestamp
        clean_name = re.sub(r'[^a-zA-Z0-9]', '_', Path(file_name).stem)
        category_code = classification["category"][:3].upper()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        return f"{category_code}_{clean_name}_{timestamp}"
    
    def _save_document_metadata(self, metadata: DocumentMetadata):
        """Save document metadata to file"""
        metadata_file = Path(METADATA_DIR) / f"{metadata.document_id}.json"
        
        try:
            with open(metadata_file, 'w') as f:
                json.dump(asdict(metadata), f, indent=2)
        except Exception as e:
            logger.error(f"Failed to save metadata for {metadata.document_id}: {e}")
    
    def _save_batch_results(self, results: Dict[str, Any]):
        """Save batch processing results"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        results_file = Path(PROCESSING_DIR) / f"batch_results_{timestamp}.json"
        
        try:
            with open(results_file, 'w') as f:
                json.dump(results, f, indent=2)
            logger.info(f"Batch results saved to {results_file}")
        except Exception as e:
            logger.error(f"Failed to save batch results: {e}")
    
    def _create_failure_result(self, file_path: str, error: str, start_time: float) -> Dict[str, Any]:
        """Create standardized failure result"""
        return {
            "success": False,
            "file_path": file_path,
            "error": error,
            "processing_time": time.time() - start_time,
            "timestamp": datetime.now().isoformat()
        }
    
    def _create_duplicate_result(self, file_path: str, existing_id: str, start_time: float) -> Dict[str, Any]:
        """Create standardized duplicate result"""
        return {
            "success": True,
            "duplicate": True,
            "file_path": file_path,
            "existing_document_id": existing_id,
            "message": "Document already processed (duplicate detected)",
            "processing_time": time.time() - start_time,
            "timestamp": datetime.now().isoformat()
        }
    
    def get_ingestion_statistics(self) -> Dict[str, Any]:
        """Get comprehensive ingestion statistics"""
        index_stats = self.index_manager.get_index_stats()
        
        return {
            "ingestion_stats": self.ingestion_stats,
            "index_stats": index_stats,
            "performance_summary": self.index_manager.performance_tracker.get_performance_summary(),
            "document_counts_by_category": self._get_category_counts(),
            "recent_processing": self._get_recent_processing_stats()
        }
    
    def _get_category_counts(self) -> Dict[str, int]:
        """Get document counts by category"""
        category_counts = {}
        
        # Read all metadata files
        metadata_dir = Path(METADATA_DIR)
        for metadata_file in metadata_dir.glob("*.json"):
            if metadata_file.name != "document_hashes.json":
                try:
                    with open(metadata_file, 'r') as f:
                        metadata = json.load(f)
                    
                    category = metadata.get("document_category", "unknown")
                    category_counts[category] = category_counts.get(category, 0) + 1
                    
                except Exception as e:
                    logger.warning(f"Could not read metadata file {metadata_file}: {e}")
        
        return category_counts
    
    def _get_recent_processing_stats(self) -> Dict[str, Any]:
        """Get recent processing statistics"""
        # This would analyze recent batch results
        # For now, return basic stats
        return {
            "last_24_hours": {
                "documents_processed": 0,
                "average_processing_time": 0,
                "success_rate": 0
            }
        }

# =============================================================================
# STREAMLIT INTERFACE
# =============================================================================

class IngestionInterface:
    """Streamlit interface for the ingestion system"""
    
    def __init__(self):
        self.engine = SmartIngestionEngine()
    
    def render(self):
        """Render the Streamlit interface"""
        st.title("🔄 Smart Document Ingestion System")
        st.markdown("**Intelligent document processing for Lumen Navigator RAG system**")
        
        # Sidebar with statistics
        with st.sidebar:
            st.header("📊 System Status")
            self._render_system_stats()
        
        # Main content tabs
        tab1, tab2, tab3, tab4 = st.tabs(["📁 Single Document", "📂 Batch Processing", "📈 Analytics", "⚙️ Configuration"])
        
        with tab1:
            self._render_single_document_tab()
        
        with tab2:
            self._render_batch_processing_tab()
        
        with tab3:
            self._render_analytics_tab()
        
        with tab4:
            self._render_configuration_tab()
    
    def _render_system_stats(self):
        """Render system statistics in sidebar"""
        try:
            stats = self.engine.get_ingestion_statistics()
            
            st.metric("Total Processed", stats["ingestion_stats"]["total_processed"])
            st.metric("Successful", stats["ingestion_stats"]["successful_ingestions"])
            st.metric("Failed", stats["ingestion_stats"]["failed_ingestions"])
            st.metric("Duplicates Found", stats["ingestion_stats"]["duplicates_found"])
            
            # Index statistics
            st.subheader("Index Status")
            for provider, info in stats["index_stats"].items():
                if info.get("status") == "active":
                    st.metric(f"{provider.title()} Index", info["document_count"])
                else:
                    st.error(f"{provider.title()}: {info.get('error', 'Unknown error')}")
        
        except Exception as e:
            st.error(f"Could not load statistics: {e}")
    
    def _render_single_document_tab(self):
        """Render single document processing tab"""
        st.header("Process Single Document")
        
        # File upload
        uploaded_file = st.file_uploader(
            "Choose a document",
            type=['pdf', 'docx', 'txt', 'html', 'md'],
            help="Supported formats: PDF, Word, Text, HTML, Markdown"
        )
        
        col1, col2 = st.columns(2)
        
        with col1:
            force_reprocess = st.checkbox(
                "Force reprocess if duplicate",
                help="Reprocess document even if it's already been ingested"
            )
        
        with col2:
            show_classification = st.checkbox(
                "Show classification details",
                value=True,
                help="Display detailed classification results"
            )
        
        if uploaded_file is not None:
            # Save uploaded file temporarily
            temp_path = Path(PROCESSING_DIR) / uploaded_file.name
            
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            if st.button("🚀 Process Document", type="primary"):
                with st.spinner("Processing document..."):
                    try:
                        result = self.engine.ingest_document(str(temp_path), force_reprocess)
                        
                        if result["success"]:
                            if result.get("duplicate"):
                                st.warning(f"📄 Document is a duplicate of: {result['existing_document_id']}")
                            else:
                                st.success("✅ Document processed successfully!")
                                
                                # Show results
                                if show_classification:
                                    self._display_processing_results(result)
                        else:
                            st.error(f"❌ Processing failed: {result['error']}")
                    
                    except Exception as e:
                        st.error(f"❌ Unexpected error: {e}")
                    
                    finally:
                        # Clean up temp file
                        if temp_path.exists():
                            temp_path.unlink()
    
    def _render_batch_processing_tab(self):
        """Render batch processing tab"""
        st.header("Batch Document Processing")
        
        # Directory input
        directory_path = st.text_input(
            "Directory Path",
            value="./documents",
            help="Path to directory containing documents to process"
        )
        
        col1, col2 = st.columns(2)
        
        with col1:
            recursive = st.checkbox("Include subdirectories", value=True)
        
        with col2:
            file_types = st.multiselect(
                "File types to process",
                options=['pdf', 'docx', 'txt', 'html', 'md'],
                default=['pdf', 'docx'],
                help="Select which file types to process"
            )
        
        if st.button("🗂️ Process Directory", type="primary"):
            if not Path(directory_path).exists():
                st.error(f"Directory not found: {directory_path}")
            else:
                file_patterns = [f"*.{ext}" for ext in file_types]
                
                with st.spinner("Processing directory..."):
                    try:
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        # This is a simplified version - you'd want to implement progress tracking
                        result = self.engine.ingest_directory(
                            directory_path, 
                            recursive=recursive,
                            file_patterns=file_patterns
                        )
                        
                        progress_bar.progress(1.0)
                        status_text.text("Processing complete!")
                        
                        # Display results
                        self._display_batch_results(result)
                    
                    except Exception as e:
                        st.error(f"Batch processing failed: {e}")
    
    def _render_analytics_tab(self):
        """Render analytics and statistics tab"""
        st.header("📈 Ingestion Analytics")
        
        try:
            stats = self.engine.get_ingestion_statistics()
            
            # Overall metrics
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric(
                    "Total Documents",
                    stats["ingestion_stats"]["total_processed"]
                )
            
            with col2:
                success_rate = 0
                total = stats["ingestion_stats"]["total_processed"]
                if total > 0:
                    success_rate = stats["ingestion_stats"]["successful_ingestions"] / total * 100
                st.metric(
                    "Success Rate",
                    f"{success_rate:.1f}%"
                )
            
            with col3:
                st.metric(
                    "Duplicates Found",
                    stats["ingestion_stats"]["duplicates_found"]
                )
            
            with col4:
                # Calculate total documents across indexes
                total_indexed = sum(
                    info.get("document_count", 0) 
                    for info in stats["index_stats"].values()
                )
                st.metric(
                    "Total Indexed",
                    total_indexed
                )
            
            # Category distribution
            st.subheader("📋 Documents by Category")
            category_counts = stats["document_counts_by_category"]
            
            if category_counts:
                # Create a simple bar chart using Streamlit
                categories = list(category_counts.keys())
                counts = list(category_counts.values())
                
                chart_data = pd.DataFrame({
                    'Category': categories,
                    'Count': counts
                })
                
                st.bar_chart(chart_data.set_index('Category'))
            else:
                st.info("No documents processed yet")
            
            # Index performance
            st.subheader("⚡ Index Performance")
            for provider, perf in stats["performance_summary"]["providers"].items():
                with st.expander(f"{provider.title()} Index Performance"):
                    pcol1, pcol2, pcol3 = st.columns(3)
                    
                    with pcol1:
                        st.metric("Performance Score", f"{perf['score']:.3f}")
                    
                    with pcol2:
                        st.metric("Avg Response Time", f"{perf['response_time']:.3f}s")
                    
                    with pcol3:
                        st.metric("Success Rate", f"{perf['success_rate']:.3f}")
        
        except Exception as e:
            st.error(f"Could not load analytics: {e}")
    
    def _render_configuration_tab(self):
        """Render configuration tab"""
        st.header("⚙️ System Configuration")
        
        # Document taxonomy
        with st.expander("📋 Document Taxonomy"):
            st.json(DOCUMENT_TAXONOMY)
        
        # Index routing
        with st.expander("🎯 Index Routing Strategy"):
            st.json(INDEX_ROUTING)
        
        # Chunking strategies
        with st.expander("✂️ Chunking Strategies"):
            st.json(CHUNKING_STRATEGIES)
        
        # Quality thresholds
        with st.expander("🎚️ Quality Thresholds"):
            st.json(QUALITY_THRESHOLDS)
        
        # System actions
        st.subheader("🔧 System Actions")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("🔄 Reload Indexes"):
                with st.spinner("Reloading indexes..."):
                    try:
                        self.engine.index_manager._load_existing_indexes()
                        st.success("Indexes reloaded successfully!")
                    except Exception as e:
                        st.error(f"Failed to reload indexes: {e}")
        
        with col2:
            if st.button("📊 Export Statistics"):
                try:
                    stats = self.engine.get_ingestion_statistics()
                    
                    # Create downloadable JSON
                    stats_json = json.dumps(stats, indent=2)
                    st.download_button(
                        label="💾 Download Statistics",
                        data=stats_json,
                        file_name=f"ingestion_stats_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                        mime="application/json"
                    )
                except Exception as e:
                    st.error(f"Failed to export statistics: {e}")
    
    def _display_processing_results(self, result: Dict[str, Any]):
        """Display detailed processing results"""
        
        # Basic info
        st.subheader("📄 Document Information")
        col1, col2 = st.columns(2)
        
        with col1:
            st.write(f"**Document ID:** {result['document_id']}")
            st.write(f"**Processing Time:** {result['processing_time']:.2f} seconds")
        
        with col2:
            st.write(f"**File:** {Path(result['file_path']).name}")
            st.write(f"**Category:** {result['classification']['category']}")
        
        # Classification details
        with st.expander("🏷️ Classification Details"):
            classification = result['classification']
            
            ccol1, ccol2 = st.columns(2)
            
            with ccol1:
                st.write(f"**Category:** {classification['category']}")
                st.write(f"**Subcategory:** {classification['subcategory']}")
                st.write(f"**Confidence:** {classification['confidence']:.3f}")
            
            with ccol2:
                st.write(f"**Regulatory Framework:** {classification['regulatory_framework']}")
                st.write(f"**Authority Level:** {classification['authority_level']}")
                st.write(f"**Compliance Criticality:** {classification['compliance_criticality']:.3f}")
        
        # Routing information
        with st.expander("🎯 Index Routing"):
            routing = result['routing']
            st.write(f"**Strategy:** {routing['strategy']}")
            st.write(f"**Reason:** {routing['reason']}")
            st.write(f"**Assigned Indexes:** {', '.join(routing['indexes'])}")
        
        # Ingestion results
        with st.expander("📥 Ingestion Results"):
            for provider, ing_result in result['ingestion_results'].items():
                if ing_result['success']:
                    st.success(f"✅ **{provider.title()}**: {ing_result['chunk_count']} chunks created")
                    st.write(f"   - Average chunk length: {ing_result['avg_chunk_length']:.0f} characters")
                else:
                    st.error(f"❌ **{provider.title()}**: {ing_result.get('error', 'Failed')}")
    
    def _display_batch_results(self, result: Dict[str, Any]):
        """Display batch processing results"""
        
        # Summary metrics
        st.subheader("📊 Batch Processing Summary")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Total Files", result['total_files'])
        
        with col2:
            st.metric("Successful", result['summary']['successful'])
        
        with col3:
            st.metric("Failed", result['summary']['failed'])
        
        with col4:
            st.metric("Duplicates", result['summary']['duplicates'])
        
        # Processing time
        total_time = result['summary']['total_processing_time']
        st.write(f"**Total Processing Time:** {total_time:.2f} seconds")
        
        if result['total_files'] > 0:
            avg_time = total_time / result['total_files']
            st.write(f"**Average Time per File:** {avg_time:.2f} seconds")
        
        # Detailed results
        with st.expander("📋 Detailed Results"):
            
            # Filter options
            filter_col1, filter_col2 = st.columns(2)
            
            with filter_col1:
                status_filter = st.selectbox(
                    "Filter by status:",
                    options=["All", "Successful", "Failed", "Duplicates"]
                )
            
            with filter_col2:
                show_errors = st.checkbox("Show error details", value=False)
            
            # Filter results
            filtered_results = result['results']
            
            if status_filter == "Successful":
                filtered_results = [r for r in filtered_results if r.get('success') and not r.get('duplicate')]
            elif status_filter == "Failed":
                filtered_results = [r for r in filtered_results if not r.get('success')]
            elif status_filter == "Duplicates":
                filtered_results = [r for r in filtered_results if r.get('duplicate')]
            
            # Display results table
            if filtered_results:
                for i, file_result in enumerate(filtered_results):
                    with st.container():
                        file_name = Path(file_result['file_path']).name
                        
                        if file_result.get('success'):
                            if file_result.get('duplicate'):
                                st.info(f"🔄 **{file_name}** - Duplicate")
                            else:
                                st.success(f"✅ **{file_name}** - Processed successfully")
                                
                                if 'classification' in file_result:
                                    classification = file_result['classification']
                                    st.write(f"   Category: {classification['category']} / {classification['subcategory']}")
                        else:
                            st.error(f"❌ **{file_name}** - Failed")
                            
                            if show_errors and 'error' in file_result:
                                st.write(f"   Error: {file_result['error']}")
                        
                        if i < len(filtered_results) - 1:
                            st.divider()
            else:
                st.info("No results match the selected filter")

# =============================================================================
# CONVENIENCE FUNCTIONS AND UTILITIES
# =============================================================================

def create_ingestion_engine() -> SmartIngestionEngine:
    """Create and return a configured SmartIngestionEngine instance"""
    return SmartIngestionEngine()

def quick_ingest(file_path: str, force_reprocess: bool = False) -> Dict[str, Any]:
    """Quick ingestion function for simple use cases"""
    engine = create_ingestion_engine()
    return engine.ingest_document(file_path, force_reprocess)

def batch_ingest(directory_path: str, recursive: bool = True, file_types: List[str] = None) -> Dict[str, Any]:
    """Quick batch ingestion function"""
    engine = create_ingestion_engine()
    
    if file_types is None:
        file_types = ['*.pdf', '*.docx', '*.txt', '*.html', '*.md']
    
    return engine.ingest_directory(directory_path, recursive, file_types)

def validate_system_setup() -> Dict[str, Any]:
    """Validate that the ingestion system is properly set up"""
    validation_results = {
        "directories": {},
        "embedding_models": {},
        "dependencies": {},
        "overall_status": "unknown"
    }
    
    # Check directories
    required_dirs = [DB_BASE_DIR, METADATA_DIR, PERFORMANCE_DIR, PROCESSING_DIR]
    
    for directory in required_dirs:
        dir_path = Path(directory)
        validation_results["directories"][directory] = {
            "exists": dir_path.exists(),
            "writable": dir_path.exists() and os.access(dir_path, os.W_OK)
        }
    
    # Check embedding models
    try:
        openai_model = OpenAIEmbeddings(model="text-embedding-3-small", show_progress_bar=False)
        validation_results["embedding_models"]["openai"] = {"status": "available", "error": None}
    except Exception as e:
        validation_results["embedding_models"]["openai"] = {"status": "error", "error": str(e)}
    
    try:
        google_model = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        validation_results["embedding_models"]["google"] = {"status": "available", "error": None}
    except Exception as e:
        validation_results["embedding_models"]["google"] = {"status": "error", "error": str(e)}
    
    # Check dependencies
    required_packages = {
        "langchain": "langchain",
        "faiss": "faiss-cpu",
        "openai": "openai",
        "google": "google-generativeai",
        "spacy": "spacy",
        "pandas": "pandas",
        "streamlit": "streamlit"
    }
    
    for package_name, import_name in required_packages.items():
        try:
            __import__(import_name.replace("-", "_"))
            validation_results["dependencies"][package_name] = {"status": "available", "error": None}
        except ImportError as e:
            validation_results["dependencies"][package_name] = {"status": "missing", "error": str(e)}
    
    # Determine overall status
    all_dirs_ok = all(
        info["exists"] and info["writable"] 
        for info in validation_results["directories"].values()
    )
    
    any_embedding_ok = any(
        info["status"] == "available" 
        for info in validation_results["embedding_models"].values()
    )
    
    critical_deps_ok = all(
        validation_results["dependencies"][dep]["status"] == "available"
        for dep in ["langchain", "faiss", "pandas", "streamlit"]
    )
    
    if all_dirs_ok and any_embedding_ok and critical_deps_ok:
        validation_results["overall_status"] = "ready"
    elif not all_dirs_ok:
        validation_results["overall_status"] = "directory_issues"
    elif not any_embedding_ok:
        validation_results["overall_status"] = "embedding_issues"
    elif not critical_deps_ok:
        validation_results["overall_status"] = "dependency_issues"
    else:
        validation_results["overall_status"] = "unknown_issues"
    
    return validation_results

def setup_system() -> bool:
    """Set up the ingestion system (create directories, etc.)"""
    try:
        # Create necessary directories
        for directory in [DB_BASE_DIR, METADATA_DIR, PERFORMANCE_DIR, PROCESSING_DIR, DOCUMENTS_DIR]:
            Path(directory).mkdir(exist_ok=True)
            logger.info(f"Created/verified directory: {directory}")
        
        # Create default configuration files
        embedding_models_config = {
            "openai": {
                "model": "text-embedding-3-small",
                "status": "active"
            },
            "google": {
                "model": "models/embedding-001", 
                "status": "active"
            }
        }
        
        config_file = Path(METADATA_DIR) / "embedding_models.json"
        if not config_file.exists():
            with open(config_file, 'w') as f:
                json.dump(embedding_models_config, f, indent=2)
            logger.info("Created embedding models configuration")
        
        logger.info("System setup completed successfully")
        return True
        
    except Exception as e:
        logger.error(f"System setup failed: {e}")
        return False

# =============================================================================
# EXAMPLE USAGE AND TESTING
# =============================================================================

if __name__ == "__main__":
    # Example usage and testing
    
    print("=" * 60)
    print("LUMEN NAVIGATOR SMART INGESTION SYSTEM")
    print("=" * 60)
    
    # Validate system setup
    print("\n1. Validating system setup...")
    validation = validate_system_setup()
    print(f"Overall Status: {validation['overall_status']}")
    
    if validation['overall_status'] != "ready":
        print("\nSetting up system...")
        if setup_system():
            print("✅ System setup completed")
        else:
            print("❌ System setup failed")
            exit(1)
    
    # Test document classification
    print("\n2. Testing document classification...")
    classifier = DocumentClassifier()
    
    test_content = """
    This document outlines the safeguarding policies for children's homes operated under 
    Ofsted regulations. All staff must follow these procedures to ensure child protection 
    and comply with statutory requirements.
    """
    
    classification = classifier.classify_document(test_content, "safeguarding_policy.pdf")
    print(f"Classification: {classification['category']} / {classification['subcategory']}")
    print(f"Regulatory Framework: {classification['regulatory_framework']}")
    print(f"Compliance Criticality: {classification['compliance_criticality']:.3f}")
    
    # Test chunking
    print("\n3. Testing smart chunking...")
    chunking_engine = SmartChunkingEngine()
    
    chunks_openai = chunking_engine.create_chunks(
        test_content, "openai", {"document_id": "test_doc", "category": "safeguarding"}
    )
    
    chunks_google = chunking_engine.create_chunks(
        test_content, "google", {"document_id": "test_doc", "category": "safeguarding"}
    )
    
    print(f"OpenAI chunks: {len(chunks_openai)}")
    print(f"Google chunks: {len(chunks_google)}")
    
    # Test full ingestion engine
    print("\n4. Testing ingestion engine...")
    engine = SmartIngestionEngine()
    
    # Create a test document
    test_dir = Path(DOCUMENTS_DIR)
    test_file = test_dir / "test_safeguarding_policy.txt"
    
    test_file.write_text("""
    SAFEGUARDING POLICY FOR CHILDREN'S HOMES
    
    This policy is designed to ensure compliance with Ofsted requirements and statutory obligations 
    for the protection of children in residential care.
    
    1. INTRODUCTION
    All children have the right to protection from harm. This policy outlines our commitment to 
    safeguarding and promoting the welfare of all children in our care.
    
    2. REPORTING PROCEDURES
    Any concerns about a child's welfare must be reported immediately to the designated safeguarding 
    officer. This includes but is not limited to:
    - Physical harm or injury
    - Emotional abuse or neglect
    - Sexual abuse or exploitation
    - Missing from care incidents
    
    3. STAFF RESPONSIBILITIES
    All staff members are required to:
    - Complete safeguarding training annually
    - Report any concerns immediately
    - Maintain appropriate professional boundaries
    - Follow all policies and procedures
    
    This policy is reviewed annually and updated in accordance with statutory guidance.
    """)
    
    print(f"Created test document: {test_file}")
    
    # Process the test document
    try:
        result = engine.ingest_document(str(test_file))
        
        if result["success"]:
            print("✅ Test document processed successfully!")
            print(f"   Document ID: {result['document_id']}")
            print(f"   Category: {result['classification']['category']}")
            print(f"   Indexes: {', '.join(result['routing']['indexes'])}")
            
            # Show ingestion results
            for provider, ing_result in result['ingestion_results'].items():
                if ing_result['success']:
                    print(f"   {provider.title()}: {ing_result['chunk_count']} chunks")
        else:
            print(f"❌ Test document processing failed: {result['error']}")
    
    except Exception as e:
        print(f"❌ Unexpected error during testing: {e}")
    
    # Get final statistics
    print("\n5. Final system statistics...")
    try:
        stats = engine.get_ingestion_statistics()
        print(f"Total processed: {stats['ingestion_stats']['total_processed']}")
        print(f"Successful: {stats['ingestion_stats']['successful_ingestions']}")
        
        for provider, info in stats['index_stats'].items():
            if info.get('status') == 'active':
                print(f"{provider.title()} index: {info['document_count']} documents")
    
    except Exception as e:
        print(f"Could not retrieve statistics: {e}")
    
    print("\n" + "=" * 60)
    print("TESTING COMPLETED")
    print("=" * 60)
    print("\nTo use the Streamlit interface, run:")
    print("streamlit run smart_ingestion_system.py")
    print("\nTo integrate with your existing app, import:")
    print("from smart_ingestion_system import SmartIngestionEngine, IngestionInterface")